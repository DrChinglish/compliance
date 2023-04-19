#--------------------------------------------------------------------------#
#   自定义的视图函数中调用的方法或类
#--------------------------------------------------------------------------#
from os import walk, path
from zipfile import ZipFile
from .key_frame import Extractor
from paddleocr import PaddleOCR, draw_ocr
from PIL import Image
import numpy as np



# 处理图片数据
class ImageProcess(object): 
    
    def __init__(self):
        self.txts = []  # 图片中文字
        self.boxes = []  # 图片中文字对应坐标
        self.image = None
        self.process_result = {}  # 敏感词
        self.keyword_box = {}  # 敏感词所在坐标
        

    def init_para(self, path):
        
        self.image = Image.open(path).convert('RGB')
        # 调整图片为统一大小，保持比例不变
        # base_width = 3360
        # w_percent = base_width / float(self.image.size[0])
        # h_size = int(float(self.image.size[1]) * float(w_percent))
        # self.image = self.image.resize((base_width, h_size), Image.ANTIALIAS)
        # 调用模型处理图片
        ocr = PaddleOCR(use_angle_cls=True,det=True, lang="ch")
        result = ocr.ocr(np.array(self.image), cls=True)
        # It seems there is something strange right here:
        # result= [[ 
        # [ [[24.0, 8.0], [277.0, 8.0], [277.0, 38.0], [24.0, 38.0]], ('MTV开发模式', 0.8464175462722778) ],
        # [ [[503.0, 103.0], [592.0, 103.0], [592.0, 121.0], [503.0, 121.0]], ('templates', 0.986018717288971) ],
        #  ...
        # [ [[118.0, 345.0], [162.0, 345.0], [162.0, 368.0], [118.0, 368.0]], ('接口', 0.9990861415863037) ],
        # [ [[591.0, 352.0], [741.0, 352.0], [741.0, 365.0], [591.0, 365.0]], ('nttps://blog.csdn.net/v', 0.8825844526290894) ]
        # ]]
        # So I modified it.(Using paddle-gpu version)-Zhu Yiding
        result =result[0]  
        self.boxes = [line[0] for line in result ]
        self.txts = [line[1][0] for line in result ]
        

    # 处理图片上的繁体字
    def process_traditional_characters(self):
        count = 0    #繁体字数量
        traditional_item = []   #繁体字
        traditional_characters = [line for line in self.txts if len(self.recongnize_traditional(line))]
        traditional_characters_box = [box for box,line in zip(self.boxes,self.txts) if len(self.recongnize_traditional(line))]
        img_ocr = draw_ocr(self.image, traditional_characters_box, font_path='./fonts/simfang.ttf')
      
        img = np.asarray(img_ocr)
        image_traditional = Image.fromarray(np.uint8(img))
        for character in traditional_characters:
            for item in character:
                if len(self.recongnize_traditional(item)):
                    traditional_item.append(item)
                    count += 1
                else:
                    pass
        
        self.process_result['traditional_characters'] = {'count':count,'traditional_item':traditional_item,'image':self.get_img_base64(image_traditional)}

        

    # 处理图片上的敏感词
    def process_sensitive_word(self):
        from paddleocr import draw_ocr
        from PIL import Image
        import numpy as np
        
        count = 0    #敏感词数量
        senstive_item = []   #敏感词
        senstive_characters = [line for line in self.txts if len(self.recongnize_sensitive(line))]
        senstive_characters_box = [box for box,line in zip(self.boxes,self.txts) if len(self.recongnize_sensitive(line))]
        img_ocr = draw_ocr(self.image, senstive_characters_box, font_path='./fonts/simfang.ttf')

        img = np.asarray(img_ocr)
        image_sensitive = Image.fromarray(np.uint8(img))
      
        for line in senstive_characters:
            
            if len(self.recongnize_sensitive(line)):
                senstive_item += self.recongnize_sensitive(line)
                count += 1
                
        self.process_result['senstive_characters'] = {'count':count,'senstive_item':senstive_item,'image':self.get_img_base64(image_sensitive)}
        


    # 处理图片上的英文
    def process_english_word(self):
        
        
        count = 0    #敏感词数量
        english_item = []   #敏感词
        senstive_characters = [line for line in self.txts if len(self.recongnize_english(line))]
        senstive_characters_box = [box for box,line in zip(self.boxes,self.txts) if len(self.recongnize_english(line))]
        img_ocr = draw_ocr(self.image, senstive_characters_box, font_path='./fonts/simfang.ttf')

        img = np.asarray(img_ocr)
        image_english = Image.fromarray(np.uint8(img))
      
        for line in senstive_characters:
            
            if len(self.recongnize_english(line)):
                english_item += self.recongnize_english(line)
                count += 1

        self.process_result['english_word'] = {'count':count,'english_item':english_item,'image':self.get_img_base64(image_english)}


    # 处理图片上的骷髅
    def process_skull(self):
        import sys
        sys.path.append("yolo")
        sys.path.append("yolo/model_data")
        
        from yolo.yolo import YOLO
        yolo = YOLO()     
        r_image= yolo.detect_image(self.image, crop = False, count=True)


        # self.get_img_base64(r_image) 将会报错：AttributeError: 'tuple' object has no attribute 'save'
        # 输出将会是(<PIL.Image.Image image mode=RGB size=1708x11233 at 0x1C5F78A0F70>, 0)，一个元组
        self.process_result['skull'] = {'image':self.get_img_base64(r_image[0])}


    # 检测血液
    def process_blood(self,path):
        import sys
        sys.path.append("image_classification")
        sys.path.append("image_classification/pytorch_image_classification")
        sys.path.append("image_classification/configs")

        from image_classification.predict import Predictor
        predictor =Predictor()   
        score = predictor.blood_predict(path)
        if score > 0.8:
            self.process_result['is_blood'] =  '检测到图片中含有血液'
        else:
            self.process_result['is_blood'] =  '检测到图片中不含有血液'

    # 健康游戏忠告
    def game_advice(self):
        from zhon.hanzi import punctuation
        import docx
        path = './media/filter/game_advice.docx'
        content=[]
        file = docx.Document(path)
        advice_boxes = []
        game_advice = 0
        advice_boxes = []
        game_advice = 0
        
        for para in file.paragraphs:
            content.append(str(para.text))
        
        image_text = " ".join(self.txts)
        for i in punctuation:
            image_text = image_text.replace(i,'')

        for i in content:
            if i in image_text:
                game_advice += 1
                for txt in self.txts:
                    if i in txt:
                        advice_boxes.append(self.boxes[self.txts.index(txt)])

        img_ocr = draw_ocr(self.image, advice_boxes)
        rate = game_advice/len(content)
        coverage_rate = format(rate, '.2%')
        img_result = Image.fromarray(np.array(img_ocr))
        # img_result.show()
        context = img_base64(img_result)
        description = ''
        if rate==0:
            description = '图片不含有游戏健康忠告内容'
        elif rate==1:
            description = '图片含有全部游戏健康忠告内容'
        else:
            description = '图片含有部分游戏健康忠告内容'


        self.process_result['game_advice'] = {'description':description,'coverage_rate':coverage_rate,'image':context}
        

        

        
        


    # PIL映像从 Django REST 框架后端获取到前端
    def get_img_base64(self,image):
        context = img_base64(image)
        # self.process_result['image'] = context
        return context


    # 判断是否有繁体字
    def recongnize_traditional(self,text):
        filter = TrantitionalCharacterFilter()
        filter.filter(text)
        res = filter.trantitional_list()
        return res    

    # 判断是否有敏感词
    def recongnize_sensitive(self,text): 
        keywords_path = 'media/filter/keywords.docx'
        filter = DFAFilter()
        filter.init_chains(keywords_path)
        filter.filter(text)
        res = filter.sensitive_list()
        return res

    # 判断是否有英文
    def recongnize_english(self,text): 
        res,_ = english(text)
        return res





# 处理文档数据
class DocProcess(object): 
    def __init__(self):
        self.txts = []  # 文档文字，按段落成list
        self.string = '' 
        self.process_result = {}  # 敏感词
 
    def init_para(self, path):
        import os
        docpath = path
        print(path)
        if os.path.splitext(path)[1] == '.doc' :
            # Preprocess doc 
            if not os.path.exists(path+'x'):
                from win32com import client as wc
                word = wc.Dispatch("Word.Application")
                doc = word.Documents.Open(path)
                doc.SaveAs(path+'x',12, False, "", True, "", False, False, False, False) 
                doc.Close()
                word.Quit()
            docpath = path+'x'    
        import docx
        file = docx.Document(docpath)
        for para in file.paragraphs:
            self.txts.append(str(para.text))
            self.string += str(para.text)
        
    # 处理文档的繁体字
    def process_traditional_characters(self):      
        count = 0    #繁体字数量
        traditional_item = []   #繁体字       
        for item in self.string:
            if len(self.recongnize_traditional(item)[0]):
                traditional_item.append(item)
                count += 1
        ret = self.recongnize_traditional(self.string)[1]
        self.process_result['traditional_characters'] = {'count':count,'traditional_item':traditional_item,'fulltext':ret}
     

    # 处理文档的敏感词
    def process_sensitive_word(self):    
        count = 0    #敏感词数量
        senstive_item = []   #敏感词             
        if len(self.recongnize_sensitive(self.string)[0]):
            senstive_item += self.recongnize_sensitive(self.string)[0]
            count = len(senstive_item)
        ret =   self.recongnize_sensitive(self.string)[1]
        self.process_result['senstive_characters'] = {'count':count,'senstive_item':senstive_item, 'fulltext':ret}
        

    # 处理文档的英文
    def process_english_word(self):
        count = 0    #敏感词数量
        english_item = []   #敏感词  
        if len(self.recongnize_english(self.string)[0]):
            english_item += self.recongnize_english(self.string)[0]
            count = len(english_item)
        ret = self.recongnize_english(self.string)[1]
        self.process_result['english_word'] = {'count':count,'english_item':english_item, 'fulltext':ret}


    # 判断是否有繁体字
    def recongnize_traditional(self,text):
        filter = TrantitionalCharacterFilter()
        filter.filter(text)
        res = filter.trantitional_list()
        res = list(set(res))
        return res,filter.ret

    # 判断是否有敏感词
    def recongnize_sensitive(self,text): 
        keywords_path = 'media/filter/keywords.docx'
        filter = DFAFilter()
        filter.init_chains(keywords_path)
        filter.filter(text)
        res = filter.sensitive_list()
        res = list(set(res))
        return res,filter.ret

    # 判断是否有英文
    def recongnize_english(self,text): 
        res,ret= english(text)
        res = list(set(res))
        return res, ret



# 处理语音数据
class SpeechProcess(object): 
    def __init__(self):
        self.txts = ""            # 语音转化的文字
        self.speech = None
        self.process_result = {}  # 敏感词
        self.keyword_time = {}    # 敏感词所时间

    @classmethod
    def get_converted_path(self,path):
        import os
        # return os.path.join(os.path.join(os.path.dirname(path),'converted_wav'),os.path.splitext(os.path.basename(path))[0],'.wav')
        return '{0}/converted_wav/{1}.wav'.format(os.path.dirname(path),os.path.splitext(os.path.basename(path))[0])

    def convert_to_wav(self,path):
        import os
        from pydub import AudioSegment
        print('Converting file...')
        converteddir = os.path.join(os.path.dirname(path),'converted_wav')
        if not os.path.exists(converteddir):
            os.makedirs(converteddir)
        
        convertedpath = self.get_converted_path(path)
        audio = AudioSegment.from_file(path)
        audio.export(convertedpath,format='wav')

    def asr(self,path,local=True):
        if local:
            from paddlespeech.cli.asr.infer import ASRExecutor
            asr = ASRExecutor()
            result = asr(audio_file=path)
            return result
        else :
            from paddlespeech.server.bin.paddlespeech_client import ASRClientExecutor
            asrclient_executor = ASRClientExecutor()
            result = asrclient_executor(
                input=path,
                server_ip = '127.0.0.1',
                port = 8090,
                sample_rate = 16000,
                lang = 'zh_cn',
                audio_format = 'wav'
            )
            return result

    def init_para(self, path,local=True):
        print('init speech process')
        from paddlespeech.cli.text.infer import TextExecutor
        print('text')
        # import importlib
        # ASRExecutor = importlib.import_module("paddlespeech.cli.asr.infer.ASRExecutor")
        # from paddlespeech.cli.asr.infer import ASRExecutor
        import os
        import soundfile as sf
        print('import')
        ext = os.path.splitext(path)[1].strip('.')
        if ext.upper() not in sf.available_formats().keys():
            if not os.path.exists(self.get_converted_path(path)):
                self.convert_to_wav(path)
            path = self.get_converted_path(path)
        self.re_sample(path)
        print('asr')
        result = self.asr(path=path,local=local)
        print('asr done')
        
        text_punc = TextExecutor()
        result = text_punc(text=result)
        self.txts = result

    #将音频的采样率改成16000,方便模型调用
    def re_sample(self, path):
        print('resampling')
        import librosa
        import soundfile as sf
        import numpy as np
        src_sig,sr = sf.read(path)                                 #path是要 输入的wav 返回 src_sig:音频数据  sr:原采样频率  
        src_sig =src_sig.T
        dst_sig = librosa.resample(np.asarray(src_sig),sr,16_000)  #resample 入参三个 音频数据 原采样频率 和目标采样频率
        dst_sig = dst_sig.T
        sf.write(path,dst_sig,16_000)                              #写出数据  参数三个 ：  目标地址  更改后的音频数据  目标采样数据
     

    # 处理敏感词
    def process_sensitive_word(self):    
        print('sensitive')
        count = 0    #敏感词数量
        senstive_item = []   #敏感词             
        if len(self.recongnize_sensitive(self.txts)[0]):
            senstive_item += self.recongnize_sensitive(self.txts)[0]
            count = len(senstive_item)
        ret =   self.recongnize_sensitive(self.txts)[1]
        self.process_result['senstive_characters'] = {'count':count,'senstive_item':senstive_item, 'fulltext':ret}
        

    # 处理英文
    def process_english_word(self):
        print('english')
        count = 0    #敏感词数量
        english_item = []   #敏感词  
        if len(self.recongnize_english(self.txts)[0]):
            english_item += self.recongnize_english(self.txts)[0]
            count = len(english_item)
        ret = self.recongnize_english(self.txts)[1]
        self.process_result['english_word'] = {'count':count,'english_item':english_item, 'fulltext':ret}


    # 判断是否有敏感词
    def recongnize_sensitive(self,text): 
        keywords_path = 'media/filter/keywords.docx'
        filter = DFAFilter()
        filter.init_chains(keywords_path)
        filter.filter(text)
        res = filter.sensitive_list()
        res = list(set(res))
        return res,filter.ret

    # 判断是否有英文
    def recongnize_english(self,text): 
        res,ret= english(text)
        res = list(set(res))
        return res, ret



# 处理视频数据
class VideoProcess(object): 
    def __init__(self):
        self.frame_path = []            # 关键帧保存路径
        self.frame_time = []            # 关键帧时间
        self.frame = []
    
    def extract_frame(self, path):
        extractor=Extractor(path)
        self.frame_path = [frame.frame_path for frame in extractor.frames]
        self.frame_time = [f'{frame.hour}:{frame.minute}:{frame.second}:{frame.millisecond}' for frame in extractor.frames]
        self.frame = [frame.id for frame in extractor.frames]



# 敏感词处理
class DFAFilter(object):
    def __init__(self):
        self.keyword_chains = {}    # 敏感词链表
        self.delimit = '\x00'       # 敏感词词尾标识
        self.whitelist=['的']       # 白名单
        self.ret = []               # 返回字符串列表，为了便于前端显示，采取形如[{'flag':0, 'text':"abc"}, {'flag':1 ,'text':"sb"}]的形式返回
    def add(self, keyword):
        keyword = keyword.lower()   # 关键词英文变为小写
        chars = keyword.strip()     # 关键字去除首尾空格和换行
        if not chars:               # 如果关键词为空直接返回
            return
        level = self.keyword_chains
        # 遍历关键字的每个字
        for i in range(len(chars)):
            # 如果这个字已经存在字符链的key中就进入其子字典
            if chars[i] in level:
                level = level[chars[i]]
            else:
                if not isinstance(level, dict):
                    break
                for j in range(i, len(chars)):
                    level[chars[j]] = {}
                    last_level, last_char = level, chars[j]
                    level = level[chars[j]]
                last_level[last_char] = {self.delimit: 0}
                break
        if i == len(chars) - 1:
            level[self.delimit] = 0
        
    def init_chains(self, path):
        import docx
        import re
        keyword_file = docx.Document(path)
        keywords = []
        for para in keyword_file.paragraphs:
            keywords.extend(re.split("[；|;]",str(para.text)))
        for keyword in keywords:
            self.add(str(keyword).strip())

    def filter(self, message):
        message = message.lower()
        start = 0
        while start < len(message):
            level = self.keyword_chains
            step_ins = 0
            for char in message[start:]:
                if char in level:
                    step_ins += 1
                    if self.delimit not in level[char] or message[start:start+step_ins] in self.whitelist: 
                        level = level[char]
                    else:
                        if start+step_ins < len(message) and (message[step_ins+start+1] in level[char]):
                            level = level[char]
                        else:
                            self.ret.append({'flag':1,'text':message[start:start+step_ins]})
                            start += step_ins - 1
                            break
                else:
                    self.ret.append({'flag':0,'text':message[start]})
                    break
            else:
                self.ret.append({'flag':0,'text':message[start]})
            start += 1
        return self.ret

    def sensitive_list(self):
        sensitive=[]
        for word in self.ret:  
            if word['flag']==1:
                sensitive.append(word['text'])
        return sensitive



# 繁体字词处理
class TrantitionalCharacterFilter(object):
    def __init__(self):
        self.ret = []  # 返回字符串列表，为了便于前端显示，采取形如[{'flag':0, 'text':"简体字"}, {'flag': 2 ,'text':"繁体字"}]的形式返回


    def filter(self,text):
        for i in text:
            if self.recongnize_traditional(i):
                self.ret.append({'flag':2 ,'text':i})
            else:
                self.ret.append({'flag':0 ,'text':i})


    def trantitional_list(self):
        trantitional=[]
        for word in self.ret:  
            if word['flag']==2:
                trantitional.append(word['text'])
        return trantitional
        

    # 判断是否有繁体字
    def recongnize_traditional(self,text):
        from zhconv import convert
        try:
            text.encode('big5hkscs')
            if convert(text, 'zh-hans')!=text:
                return True
            else:
                return False
        except:
            if convert(text, 'zh-hans')!=text:
                return True
            else:
                return False
      

# 英文处理
def english(string):
    import re
    adressRegex = re.compile(r'''(
                [a-zA-Z]+
                    )''', re.VERBOSE)
    res = adressRegex.findall(string) 
    whitelist =  ['PC','3D','2D','H5','VR','AR','HD','Q','K']
    for i in whitelist:
            if i in res:
                res.pop(res.index(i))

     # ret = []# 返回字符串列表，为了便于前端显示，采取形如[{'flag':0, 'text':"非英文"}, {'flag':3 ,'text':"english"}]的形式返回
    # for idenx,item in enumerate(re.split("[ ,.，。]",string)):
    #     if item in res:
    #         ret.append({'flag':3 ,'text':item})
    #     else:
    #         ret.append({'flag':0 ,'text':item})
    import string as s
    eng_word=''
    word=''
    lasttype = 0 #标识上一次检测到的字符类型 0：初始，1：其他文本，-1：英文
    wordlist=[]
    # print(string)
    for ch in string:
        if ch not in s.ascii_lowercase + s.ascii_uppercase:
            if lasttype * 1 == -1:
                # 检测到了新的类型
                if eng_word in whitelist:
                    wordlist.append({'flag':0 ,'text':eng_word})
                else:
                    wordlist.append({'flag':3 ,'text':eng_word})
                eng_word = ''
            # Not english
            word+=ch
            lasttype=1
        else:
            if lasttype * 1 == -1:
                # 检测到了新的类型
                wordlist.append({'flag':0 ,'text':word})
                word = ''
            eng_word+=ch
            lasttype = -1
    if lasttype == 1:
        if len(word)>0:
            wordlist.append({'flag':0 ,'text':word})
    elif lasttype == -1:
        if len(eng_word)>0:
            wordlist.append({'flag':3 ,'text':eng_word})
    return res, wordlist





# 链接数据库
class DBConnection(object):
    
    def __init__(self,user,pwd,dbname,tablename):
        self.user = user
        self.pwd = pwd
        self.dbname = dbname
        self.tablename = tablename
        self.formheader=[]
        
    def conn(self):
        import pymysql
        database = pymysql.connect(host='localhost',
                       port=3306,
                       user=self.user,
                       passwd=self.pwd,                     
                       db=self.dbname,
                       charset = 'utf8')

        cursor = database.cursor()
        return cursor

    def get_data(self):
        sql = "select * from {}".format(self.tablename)
        cursor = self.conn() 
        cursor.execute(sql)
        self.formheader=[]
        desc = cursor.description
        for field in desc:
            self.formheader.append(field[0])
        # print(self.formheader)
        ret = cursor.fetchall()
        
        return ret
    
    def get_tables(self,db_name):
        import pymysql
        sql = 'show tables;'
        cursor = pymysql.connect(host='localhost',
                       port=3306,
                       user=self.user,
                       passwd=self.pwd,
                       db=db_name,
                       charset = 'utf8').cursor()
        cursor.execute(sql)
        ret = cursor.fetchall()
        ret_list = [r[0] for r in ret]
        cursor.close()
        print(ret)
        return ret_list
        

    def get_databases(self):
        import pymysql
        sql = 'show databases;'
        cursor = pymysql.connect(host='localhost',
                       port=3306,
                       user=self.user,
                       passwd=self.pwd,
                       charset = 'utf8').cursor()
        cursor.execute(sql)
        ret = cursor.fetchall()
        ret_list = [r[0] for r in ret]
        cursor.close()
        print(ret)
        return ret_list


# 获取文件夹下的各类文件,接受两个参数，第一个为文件根目录，第二个为要获取的文件类型
def get_file(rootdir, type):
    img_data = []
    doc_data = []
    audio_data = []
    video_data = []
    for root, dirs, files in walk(rootdir, topdown=True):
        for name in files:
            pre, ending = path.splitext(name)
            if ending == ".jpg" or ending == ".jepg" or ending == ".png":
                img_data.append(name)
            if ending == ".docx" or ending == ".txt" or ending == ".doc":
                doc_data.append(name)
            if ending == ".wav" or ending == ".mp3" or ending == ".wma":
                audio_data.append(name)
            if ending == ".mp4" or ending == ".m4v" or ending == ".avi":
                video_data.append(name)
            else:  
                continue
        break
    if type == "image":
        return img_data
    if type == "doc":
        return doc_data
    if type == "audio":
        return audio_data
    if type == "video":
        return video_data



# 解压用户上传的zip文件包
def unzip_file(zip_file: ZipFile):
    name_to_info = zip_file.NameToInfo
    for name, info in name_to_info.copy().items():
        real_name = name.encode('cp437').decode('gbk')  # 解决中文乱码问题
        if real_name != name:
            info.filename = real_name
            del name_to_info[name]
            name_to_info[real_name] = info
    return zip_file


# PIL映像从 Django REST 框架后端获取到前端
def img_base64(image):
    import io
    import base64
    im_io = io.BytesIO()
    print(image)
    image.save(im_io, 'png', quality=70)
    im_io.seek(0)
    im_io_png = base64.b64encode(im_io.getvalue())
    context = im_io_png.decode('UTF-8')
    return context


# 将文件转化base64编码
def file_base64(filepath):
    import base64
    with open(filepath, 'rb') as f1:
        base64_str = base64.b64encode(f1.read())  
        src = base64_str.decode('utf-8')  
    return src