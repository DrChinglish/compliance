from urllib.parse import urlparse
import requests
from .util import *
from .repattern import *
import numpy as np
from pathlib import Path
from PIL import Image, ImageDraw
from paddleocr import PaddleOCR
from mediapipe.python.solutions.face_detection import FaceDetection
from mediapipe.python.solutions.drawing_utils import _normalized_to_pixel_coordinates
from mediapipe.python.solutions.hands import Hands
import cv2

# from myOCR.quan.predict.predict import MyOCR
from platformapi.utils.law import get_law_list
from django.conf import settings
# 数据处理
class DataProcess(object):
    def __init__(self,path,path_list,modules=settings.WEBSCAN_DEFAULT_SCAN_MODULES,options=settings.WEBSCAN_DEFAULT_OCR_OPTIONS):
        self.path = path
        self.path_list = path_list # 文件本地的地址和网址的映射关系
        self.txts = []  # 文本内容
        self.string = ''
        self.boxes = [] # 图片中文字边框
        self.file_type = None
        self.image = None
        self.speech = None
        self.sensitive_information = {}
        self.scan_modules = modules
        self.options = options
        import logging,time
        self.logger = logging.getLogger(__name__)

        #   {'specified_identity':{'passport':[],'IDNumber':[],'officer':[],'HM_pass':[],'carnum':[]},
        #                          'bioinformation':{'face':[],'fingerprint':[],'iris':[],'Eye_pattern':[],'auricle':[]},
        #                          'financial_account':{'debit':[],'fund_account':[],'Alipay_account':[]},
        #                          'healthcare':{'birth_information':[],'hospital_records':[]},
        #                          'track':{'Latitude_longitude':[],'adress':[]},
        #                          'juveniles_information':{'juveniles':[]},
        #                          'authentication_information':{'password':[],'auth_code':[]},
        #                          'other':{'sex':[],'nation':[]},
        #                      } 
        
    def init_para(self, path):
        file_path = Path(path)
        if convert_type(file_path) == 'text':
            self.file_type = 'text'
            import docx
            file = docx.Document(file_path)          
            for para in file.paragraphs:
                self.txts.append([str(para.text),0])
                self.string = self.string+str(para.text)
        elif convert_type(file_path) == 'audio':
            self.file_type = 'audio'
            from paddlespeech.cli.asr.infer import ASRExecutor
            from paddlespeech.cli.text.infer import TextExecutor
            asr = ASRExecutor()
            self.re_sample(path)
            result = asr(audio_file=path)

            text_punc = TextExecutor()
            result = text_punc(text=result)
            self.txts = [result, 0]
            self.string = ''.join((str(x[0]) for x in self.txts))
        elif convert_type(file_path) == 'image':
            name = os.path.split(path)[1]
            from itertools import product
            self.file_type = 'image'
            self.image = Image.open(path).convert('RGB')
            from django.conf import settings
            import requests,json
            headers = {"Content-type": "application/json"}
            from .util import cv2_to_base64
            if 'watermark_remove' in self.options:
                img = Image.open(str(path))
                width, height = img.size
                print(width,height)
                for pos in product(range(width), range(height)):
                    if sum(img.getpixel(pos)[:3]) > 600:
                        img.putpixel(pos, (255,255,255))
                if not os.path.exists('media/ocr_test/temp'):
                    os.makedirs('media/ocr_test/temp')
                img.save(f'media/ocr_test/temp/{name}')
                img = open(f'media/ocr_test/temp/{name}', 'rb').read()
            else:
                img = open(str(path), 'rb').read()
            data = {'images': [cv2_to_base64(img)]}
            url = settings.PADDLE_OCR_HOST
            r = requests.post(
            url=url, headers=headers, data=json.dumps(data),proxies={
                'http':'',
                'https':''
            })    
            result = r.json()["results"][0]

            # # ocr = PaddleOCR(use_angle_cls=True,det=True, lang="ch")
            # # result = ocr.ocr(np.array(self.image), cls=True)
            # # result = result[0]  
            # # self.boxes = [line[0] for line in result ]
            # # self.txts = [[line[1][0], 0] for line in result ]
            # # self.string = ''.join((str(x[0]) for x in self.txts))
            # ocr = MyOCR()
            # result=ocr.predict (path)
            # for item in result:
            #     self.txts.append([item['transcription'],0])
            #     self.boxes.append(item['points'])

            # 根据ppocr不同版本，返回格式可能有区别，需要注意
            self.boxes = [res['text_region'] for res in result]
            self.txts = [res['text'] for res in result ]
            print(f'识别出的文本内容：{self.txts}')
            self.string = ''.join((str(x)+" " for x in self.txts))
            print(f'识别出的文本字符串：{self.string}')
            if 'watermark_remove' in self.options:
                os.remove(f'media/ocr_test/temp/{name}')
        else:
            pass
        
    # 将音频的采样率改成16000,方便模型调用
    def re_sample(self, path):
        import librosa
        import soundfile as sf
        import numpy as np
        src_sig,sr = sf.read(path)                                
        src_sig =src_sig.T
        dst_sig = librosa.resample(np.asarray(src_sig),sr,16_000) 
        dst_sig = dst_sig.T
        sf.write(path,dst_sig,16_000)   



    # 查找用户上传数据中的敏感信息
    def search_risk(self): 
        if self.file_type in ['text', 'audio', 'image']:
            data = {self.path_list[self.path]:self.string}
            riskdata = SearchRiskdata(data)
            res = riskdata.miti_process()
            self.logger.info(f'从 {self.path} 中识别出的文字------->{res}')
            # print(f'图片中的文字------->{res}')
            ret = riskdata.ret
            self.sensitive_information = merge_dict(res,self.sensitive_information)

            # for item in search_riskdata(self.string):
            #     for type in self.sensitive_information.keys():
            #         if item[0] in self.sensitive_information[type].keys():
            #             self.sensitive_information[type][item[0]].append(item[3])
            #     for i,txt in enumerate(self.txts):
            #         if item[3] in txt[0]:
            #             self.txts[i][1] = 1

    # 图片中人脸检测
    def face_detect(self,path):
        if self.file_type =='image':
            path = str(Path(path))
            face_det = FaceDet(min_detection_confidence=0.5, model_selection=1)
            bboxes, _=face_det.detect_all_faces(image=cv2.cvtColor(cv2.imread(path), cv2.COLOR_RGB2BGR))
            if len(bboxes)>0:   
                law_list = get_law_list('high')
                content_list = ['图片中含人脸信息',self.path_list[path]]
                # self.sensitive_information['bioinformation']['face'].append(path)
                self.maching_law('人脸',content_list,law_list)
        return
    
    # 图片中暴露检测
    def baolu_detect(self,path):
        if self.file_type =='image':
            path = str(Path(path))
            baolu_det = BaoluDet(path)
            res = baolu_det.detect()
            if res:
                law_list = get_law_list('high')
                content_list = ['图片中含黄色敏感信息',self.path_list[path]]
                # self.sensitive_information['bioinformation']['face'].append(path)
                self.maching_law('黄色图片',content_list,law_list)
        return
    
    # 图片手指纹检测
    def fingerprint_detect(self,path):
        if self.file_type =='image':
            path = str(Path(path))
            hand_det = HandDet()
            keys=hand_det.detect_all_hands(image=cv2.cvtColor(cv2.imread(path), cv2.COLOR_RGB2BGR))
            if len(keys)>0:
                # self.sensitive_information['bioinformation']['fingerprint'].append(path)
                law_list = get_law_list('high')
                content_list = ['图片中含指纹信息',self.path_list[path]]
                self.maching_law('指纹',content_list,law_list)
        return
    
    # 图片中的其他生物信息
    def bioinfo_detect(self,path):
        if self.file_type == 'image':
            bioinfo_det = BioinfoDet() 
            path = str(Path(path))
            ret = bioinfo_det.bio_detect(onnx_path="media/files/platformapi/onnx_server/best.onnx",imgpath=path,show=True)
            if len(ret)>0:
                for item in ret:
                    if item["classes"] == 'iris':
                        # self.sensitive_information['bioinformation']['iris'].append(path)
                        law_list = get_law_list('high')
                        content_list = ['图片中含虹膜信息',self.path_list[path]]
                        self.maching_law('虹膜',content_list,law_list)
                    if item["classes"] == 'auricle':
                        # self.sensitive_information['bioinformation']['auricle'].append(path)
                        law_list = get_law_list('high')
                        content_list = ['图片中含耳廓信息',self.path_list[path]]
                        self.maching_law('耳廓',content_list,law_list)
                    return
            
    def maching_law(self,category,content_list,law_list):
        old_match = category in self.sensitive_information['overview']['high']
        if old_match:
            self.sensitive_information['overview']['high'][category] += 1
            self.sensitive_information['detail']['high'][category]['content_list'].append(content_list)
        else:
            self.sensitive_information['overview']['high'][category] = 1
            self.sensitive_information['detail']['high'][category]={'content_list':[content_list],'law_list':law_list}

    # 文本数据脱敏
    def des_txts(self):
        if self.file_type in ['text', 'audio', 'image']:
            riskdata = SearchRiskdata(self.string)
            res = riskdata.miti_process()
            ret = riskdata.ret
            # for item in search_riskdata(self.string):
            for item in ret:
                for i,para in enumerate(self.txts):
                    if item[3] in para[0]:
                        self.txts[i][0] = self.txts[i][0].replace(item[3], eval(f'des_{item[0]}')(item[3]))

    #开始处理
    def process(self):
        if isinstance(self.path, str):
            self.one_process(self.path)
        else:
            self.miti_process()

    
    #处理单个文件
    def one_process(self,path):
        
        self.init_para(path)
        if '文本类型敏感信息检测' in self.scan_modules:
            self.search_risk() 
        if '人脸信息检测' in self.scan_modules:
            self.face_detect(path)
        if '暴露信息检测' in self.scan_modules:
            self.baolu_detect(path)
        if '指纹信息检测' in self.scan_modules:
            self.fingerprint_detect(path) 
        if '生物识别信息检测' in self.scan_modules:
            self.bioinfo_detect(path)


    # # 多线程处理文件
    # def miti_process(self):
    #     all_links = self.path()
    #     # 多线程处理
    #     threads = []
    #     for link in all_links:
    #         print(link)
    #         thread = threading.Thread(target=self.one_process, args=(link,))
    #         threads.append(thread)
    #         thread.start()

    #     # 等待所有线程结束
    #     for thread in threads:
    #         thread.join()


    # 图片数据脱敏，将敏感数据部分打马赛克             
    def des_iamge(self):
        if self.file_type == 'image':
            img_array = np.array(self.image).astype(int) 
            pixel = 40
            for i, item in enumerate(self.boxes):
                if self.txts[i][1] == 1:
                    for x in range(int(item[0][1]), int(item[2][1])-20, pixel):
                        for y in range(int(item[0][0]), int(item[2][0])-20, pixel):
                            img_array[x:x + pixel, y:y + pixel] = img_array[x +(pixel // 2)][y + (pixel // 2)]
            self.image = Image.fromarray(img_array.astype("uint8"))
        else:
            return

    # 保存处理后的数据到指定的路径
    def save_des_data(self, save_path):
        import docx
        if self.file_type == 'text':
            doc = docx.Document()
            for para in self.txts:
                par = doc.add_paragraph('')
                run_ = par.add_run(para[0])
            doc.save(save_path)
        elif self.file_type == 'image':
            img_array = np.array(self.image).astype(int)
            image = Image.fromarray(img_array.astype("uint8"))
            image.save(save_path)
        else:
            return        




# 多线程处理图片
def process_img(res, path_list, modules=settings.WEBSCAN_DEFAULT_SCAN_MODULES,options=settings.WEBSCAN_DEFAULT_OCR_OPTIONS):
    # 定义一个锁对象
    lock = threading.Lock()
    import logging
    logger = logging.getLogger(__name__)
    
    def worker(path,path_list):
        nonlocal res
        try:
            dataprocesser = DataProcess(path,path_list,modules=modules,options=options)
            # dataprocesser.sensitive_information = res
            logger.info(f"Processing image file:{path}.")
            dataprocesser.process()
            
            # 获取锁
            lock.acquire()
            try:
                res =   merge_dict(res,dataprocesser.sensitive_information)
            finally:
                # 释放锁
                lock.release()
        except Exception as e:
            logger.error("Invalid link: \"{0}\", traceback: {1}".format(path,e))
            
            print("Invalid link: \"{}\"".format(path))
    
    # 多线程处理
    threads = []
    

    for link in path_list:
        logger.info(f'处理图片------>>{link}')
        print(f'处理图片------>>{link}')
        thread = threading.Thread(target=worker, args=(link,path_list))
        threads.append(thread)
        thread.start()

    # 等待所有线程结束
    for thread in threads:
        thread.join()

    return res




from .resnet import resnet50
import torch
from torchvision import transforms
from PIL import Image
# 暴露图片检测
class BaoluDet:
    
    def __init__(self,path):
        self.model = "media/para/baolu_detectNet.pth"
        self.path = path
 

    def detect(self):
        img = Image.open(self.path)
        img = img.convert('RGB')
        transf1 = transforms.Resize((224,224))
        img = transf1(img)
        transf = transforms.ToTensor()
        img = transf(img).unsqueeze(0) 
        
    
        model = resnet50()
        
        model.cuda()
        
        model.load_state_dict(torch.load(self.model))
        model.eval()
                    

        img = img.cuda()
        output = model(img)
        y = output.detach().cpu().numpy()
        result = False
        if(y>=0.5):
            result = True
        return result
    
        






# 人脸检测模型
class FaceDet(FaceDetection):

    def detect_all_faces(self, image):
        image_cols, image_rows, _=image.shape
        results=self.process(image=image)

        bboxes=[]
        keypoints=[]
        if not results.detections:
            return bboxes, keypoints
        for detection in results.detections:
            location = detection.location_data
            if not location.HasField('relative_bounding_box'):
                continue
            relative_bounding_box = location.relative_bounding_box

            x_min, y_min = _normalized_to_pixel_coordinates(
                relative_bounding_box.xmin, 
                relative_bounding_box.ymin, image_cols, image_rows)

            x_max, y_max = _normalized_to_pixel_coordinates(
                relative_bounding_box.xmin + relative_bounding_box.width,
                relative_bounding_box.ymin + relative_bounding_box.height, image_cols, image_rows)
            bboxes.append((x_min, y_min, x_max, y_max))
            
            keypoints.append([_normalized_to_pixel_coordinates(
                        keypoint.x, keypoint.y, image_cols, image_rows
                        ) for keypoint in location.relative_keypoints])
        
        return bboxes, keypoints


# 手指检测模型
class HandDet(Hands):

    def detect_all_hands(self, image):
        
        image_cols, image_rows, _=image.shape
        results=self.process(image=image)
        
        keypoints=[]
        if not results.multi_hand_landmarks:
            return keypoints

        for landmark_list in results.multi_hand_landmarks:
            if not landmark_list:
                continue
            hand_landmarks = landmark_list.landmark
            keypoints.append([(i, round(landmark.x * image_rows), round(landmark.y * image_cols))
                                for i, landmark in enumerate(hand_landmarks)])

        return keypoints 
    


# 生物信息检测模型(onnx)
from .detect_onnx_server.operation import YOLO
class BioinfoDet():

    def bio_detect(self,onnx_path,imgpath='',show=True):
        '''
        检测目标，返回目标所在坐标如：
        {'crop': [57, 390, 207, 882], 'classes': 'person'},...]
        :param onnx_path:onnx模型路径
        :param img:检测用的图片
        :param show:是否展示
        :return:
        '''
        yolo = YOLO(onnx_path=onnx_path)
        # 结果
        det_obj = yolo.decect(imgpath)

        # 画框框
        if show:
            img = Image.open(imgpath)
            draw = ImageDraw.Draw(img)

            for i in range(len(det_obj)):
                draw.rectangle(det_obj[i]['crop'],width=3)
            # img.show()  # 展示
        return det_obj




    

import asyncio
import aiohttp
import os
import re
import uuid
from urllib.parse import urlparse
from bs4 import BeautifulSoup

class ImgScraper:
    def __init__(self, urls, img,  max_concurrency=5):
        self.urls = urls
        self.img = img
        self.max_concurrency = max_concurrency
        

    async def download_image(self, link, url):
        # 获取图片名称
        img_name = os.path.basename(link)

        # 发送请求并下载图片
        async with aiohttp.ClientSession(connector=aiohttp.TCPConnector(limit=self.max_concurrency)) as session:
            try:
                async with session.get(link, timeout=1000) as response:
                    try:
                        img_content = await response.read()
                    except:
                        print(f"1-->>Invalid link: {url}")
                        return
                     # 图片大小过滤
                    if len(img_content) < 25000:
                        print(f'{img_name} 图片小于 10000 字节，不下载')
            # except  aiohttp.ClientError or asyncio.TimeoutError or asyncio.RuntimeError:
            except:
                print(f"1-->>Invalid link: {url}")
                return


        parsed_url = urlparse(url)
        # domain = parsed_url.netloc.split(':')[0]
        domain = 'www.myouhome.com'
        base = "media/files/webscan/{0}/".format(domain)
        if not os.path.exists(base):
            # 不存在则创建
            os.makedirs(base)
        path = base+img_name
        # path = base+"{0}.{1}".format(uuid.uuid4(),img_name.split('.')[1])
        with open(path, 'wb') as f:
            f.write(img_content)
        print(f'----------->{link}\n{img_name} 下载完成')
        self.img[path] = link




    async def spider(self, url):
        # 获取网站 HTML 内容
        async with aiohttp.ClientSession(connector=aiohttp.TCPConnector(limit=self.max_concurrency)) as session:
            try:
                async with session.get(url, timeout=1000) as response:
                    html_content = await response.read()    
            # except  aiohttp.ClientError or asyncio.TimeoutError:
            except:
                print(f"Failed to get content from {url}")
                return

        # 解析 HTML 内容
        soup = BeautifulSoup(html_content, 'html.parser')

        # 查找网站上所有的图片链接
        img_links = []
        for img in soup.find_all('img'):
            if img.get('src') and not img.get('zoomfile') and not img.get('file'):
                img_link = img.get('src')
            else:
                img_link = img_link = img.get('file')
            # 解析图片链接，判断是否缺少协议
            parsed_link = urlparse(img_link)
            if not parsed_link.scheme:
                img_link = f'{urlparse(url).scheme}://{urlparse(url).netloc}/{img_link}'

            # 判断是否为图片链接
            if re.search(r'\.(jpg|jpeg|png|bmp)$', img_link):
                img_links.append(img_link)

        # 异步下载图片
        tasks = []
        for link in set(img_links):
            print(f'-------------->>>{link}')
            task = asyncio.ensure_future(self.download_image(link, url))
            tasks.append(task)
        await asyncio.gather(*tasks)

    async def run(self):
        # 创建多个进程并发执行爬虫任务
        tasks = []
        async with aiohttp.ClientSession() as session:
            for url in self.urls:
                tasks.append(asyncio.create_task(self.spider(url)))
            await asyncio.gather(*tasks)

    def start(self):
        loop = asyncio.get_event_loop()
        task = loop.create_task(self.run())
        loop.run_until_complete(task)
        print(self.img)

        


# 扫描网页
import os
import threading
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from multiprocess import Process
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse

class WebScraper:
    def __init__(self,url=None,logparser=None,logfile=None):
        self.logparser = logparser # 处理爬虫日志文件
        self.logfile = logfile # 爬虫日志文件地址
    
        self.ignoring_size = 10000 # 忽略图片文件的大小阈值

        self.url = url
        self.filtered_text = {}
        self.img = {}

        import logging,time
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(level = logging.INFO)
        self.loghandler = logging.FileHandler(
            f"media/files/webscan/running_log/log-web-{time.asctime().replace(' ','-')}.log")
        self.loghandler.setLevel(logging.INFO)
        formatter = logging.Formatter('%(asctime)s - %(name)s [%(levelname)s] : %(message)s')
        self.loghandler.setFormatter(formatter)
        self.logger.addHandler(self.loghandler)
        self.logger.info('Logger initiated. ')

        

    def get_links(self):
        if self.logfile and self.logfile:
            print('Using log file...')
            return self.logparser(self.logfile)
        else:
            # 获取网站的所有链接
            r = requests.get(self.url)
            soup = BeautifulSoup(r.content, 'html.parser')

            links = []
            for link in soup.find_all(href=True):
                href = link.get('href')
                if href:
                    links.append(href)

            for img in soup.find_all('img', src=True):
                src = img.get('src')
                if src:
                    links.append(src)

            for a in soup.find_all('a', {'data-href': True}):
                data_href = a.get('data-href')
                if data_href:
                    links.append(data_href)

            parsed_url = urlparse(self.url)
            base_url = '{}://{}'.format(parsed_url.scheme, parsed_url.netloc)

            absolute_links = []
            for link in links:
                absolute_link = urljoin(base_url, link)
                if not absolute_link.endswith(('.js', '.css')):
                    absolute_links.append(absolute_link)

            return list(set(absolute_links))
    
    

    def get_data(self,url,params=None):
        # 下载url对应资源，可能是图片/html等
        try:
            res = requests.get(url=url,params=params,proxies={
                'http:':"",
                'https':""
            })
        except Exception as e:
            self.logger.error("Invalid link: \"{0}\" while trying to download resource. Trace back: \n {1}".format(url,e))
            # print("Invalid link: \"{}\"".format(url))
            return
        # print(res.text,res.headers,res.cookies.items())
        try:
            content_type,content_format,_ = self.get_response_type(res.headers['Content-Type'])
        except KeyError as e:
            self.logger.error(f"处理URL:{url}时，请求头中未发现“Content-Type”属性，请求头如下：{res.headers}")
            return
        # print(res_type)
        file_handler={
            'image':self.download_image,
            'text':self.resolve_text,
        }
        def unknown_type(type,url,content=None,format=None):
            self.logger.error('Unsupported Content-Type: {}'.format(type))
            # print('Unsupported Content-Type: {}'.format(type))
            return None
        handler = file_handler.get(content_type,unknown_type)
        if handler:
            res=handler(url=url,content=res.content,format=content_format,type=content_type) #-----ttt-------
        else:
            return
        if isinstance(res, list):
            self.filtered_text[url]=''.join(res)
        elif isinstance(res, tuple):
            self.img[res[0]]=res[1]
        else:
            return


    

    def miti_process(self):
        all_links = self.get_links()
        parsed_url = urlparse(all_links[0])
        domain = parsed_url.netloc.split(':')[0]
        self.url = domain
        # 多线程处理
        threads = []
        for link in all_links:
            self.logger.info('Starting thread for link \"{}\".'.format(link))
            # print(link)
            thread = threading.Thread(target=self.get_data, args=(link,))
            threads.append(thread)
            thread.start()

        # 等待所有线程结束
        for thread in threads:
            thread.join()

    
    def get_response_type(self,contentType:str):
        # 分析资源类型
        format_map ={
            'png':'image'
        } # 某些网站响应标头不规范的情况下适用例外
        list = contentType.split(';')
        type = list[0]
        param = list[1:] if len(list)>1 else None
        type_list = type.split('/')
        if len(type_list) == 2: # 标准的情况
            content_type = type_list[0]
            content_format = type_list[1]
        elif len(type_list) == 1 and type_list[0] in format_map.keys():
            content_type = format_map[type_list[0]]
            content_format = type_list[0]
        else:
            print("Unrecognized Content-Type: {}".format(contentType))
            return None
        # print(list,type_list)
        return content_type,content_format,param


    def save_as_file(self,content,path,url):
        # 将资源保存为文件
        try:
            self.logger.info(f'{url} --->正在保存')
            file = open(path,'wb')
            file.write(content)
            file.close()
            self.logger.info(f'{url} --->保存成功')
            # print(f'{url}--->保存成功')

            return path,url
        except:
            self.logger.error("Error while saving file \"{}\"".format(path))
            # print("Error while saving file \"{}\"".format(path))
            return None


    def download_image(self,url,content,format,type=None):
        # 保存图片
        import uuid
        # base = "imgs/"    
        base = "media/files/webscan/saved_files"
        parsed_url = urlparse(url)
        # print(parsed_url)
        # print(self.url,parsed_url.netloc,parsed_url.netloc.split(':'))
        domain = parsed_url.netloc.split(':')[0]
        print(domain)
        #path = "{0}{1}.{2}".format(base,domain,format)
        path = "{0}/{1}".format(base,domain) 
        
        if not os.path.exists(path):
            # 不存在则创建
            os.makedirs(path)
        path = "{0}/{1}/{2}.{3}".format(base,domain,uuid.uuid4(),format)
        print(path)
        # 图片大小过滤
        if len(content) < 30000:
            self.logger.info(f"下载图片:{url}，由于图片大小小于{self.ignoring_size}Bytes，不保存到本地")
            # print('图片小于 10000 字节，不下载')
            return      
        return self.save_as_file(content,path,url)
    


    def resolve_text(self,url,content,format:str,type=None):
        # 解析文本类型资源
        if format.lower() == 'html':
            return self.resolve_html(content)
        else :
            self.logger.error(f"Unsupported format while processing text file from {url}: {format}")
            # print("Unsupported format")
            return None

    def resolve_html(self,content):
        # 提取html内文本信息
        from lxml import etree 
        root_elem = etree.HTML(content)
        # xpath_text = "//descendant::text()"
        # xpath_text = "(//p)//text()"
        xpath_text = "(//a | //p | //li | //span | //h1 | //h2 | //h3 | //h4 | //h5 | //h6 | //strong)//text()"
        text_list = root_elem.xpath(xpath_text)
        # print(text_list)
        import re
        patt = re.compile(r"\S+")
        # print(patt.search(" 2 "))
        filtered_list = list(filter(lambda str: patt.search(str)!=None,text_list))
        # print("\n",filtered_list,len(text_list),len(filtered_list))
        return filtered_list


# 解压用户上传的zip文件包
from zipfile import ZipFile
def unzip_file(zip_file: ZipFile):
    name_to_info = zip_file.NameToInfo
    for name, info in name_to_info.copy().items():
        real_name = name.encode('cp437').decode('gbk')  # 解决中文乱码问题
        if real_name != name:
            info.filename = real_name
            del name_to_info[name]
            name_to_info[real_name] = info
    return zip_file                


# 链接数据库
class DBConnection(object):
    
    def __init__(self,user,pwd,dbname='',tablename='',dbtype='mysql',host='localhost'):
        self.host = host
        self.user = user
        self.pwd = pwd
        self.dbname = dbname
        self.tablename = tablename
        self.formheader=[]
        self.dbtype = dbtype
    
    def db_list(self):
        sql_my = "show databases;"
        sql_ms = "select name from sysdatabases;"
        sql = sql_my if self.dbtype=='mysql' else sql_ms
        cursor = self.conn()
        cursor.execute(sql)
        res = cursor.fetchall()
        rl = [r[0] for r in res]
        return rl

    def table_list(self):
        sql_ms = "select name from sysobjects where xtype='U';"
        sql_my = "show tables;"
        sql = sql_my if self.dbtype=='mysql' else sql_ms
        
        cursor = self.conn_db()
        cursor.execute(sql)
        res = cursor.fetchall()
        rl = [r[0] for r in res]
        return rl

    def conn(self):
        import pymysql,pymssql
        print(self.host)
        if self.dbtype == 'mssql': # Sql Server
            # port = self.port
            # if not self.port:
            #     port = 1433
            database = pymssql.connect(host=self.host,
                                       port=1433,
                                       user=self.user,
                                       password=self.pwd
                                      
                                       )
        else: #Mysql 
            database = pymysql.connect(host=self.host,
                        port=13306,
                        user=self.user,
                        passwd=self.pwd,                     
                        db=self.dbname,
                        charset = 'utf8')

        cursor = database.cursor()
        return cursor
    
    def conn_db(self):
        import pymysql,pymssql
        if self.dbtype == 'mssql': # Sql Server
            # port = self.port
            # if not self.port:
            #     port = 1433
            print(self.user,self.pwd,self.dbname,self.host)
            database = pymssql.connect(host=self.host,
                                       port=1433,
                                       user=self.user,
                                       password=self.pwd,
                                       database=self.dbname
                                       )
        else: #Mysql 
            database = pymysql.connect(host=self.host,
                        port=13306,
                        user=self.user,
                        passwd=self.pwd,                     
                        db=self.dbname,
                        charset = 'utf8')

        cursor = database.cursor()
        return cursor

    def get_data(self):
        sql = "select * from {}".format(self.tablename)
        cursor = self.conn_db()
        cursor.execute(sql)
        self.formheader=[]
        desc = cursor.description
        for field in desc:
            self.formheader.append(field[0])
        # print(self.formheader)
        # ret = cursor.fetchall()
        ret = cursor.fetchmany(size=100)
        cursor.close()
        return ret