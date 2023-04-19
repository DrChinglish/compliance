from http.client import HTTPResponse
import imp
import sys
from turtle import title
from unittest.mock import patch

sys.path.append("./../env/Lib/site-packages/bert")
sys.path.append("./../env/Lib/site-packages/bert/NER")
# sys.path.append("./../env/Lib/site-packages")
#import NER.predict as bert
# import bert
from .DBConnection import DBConnection
import datetime
from multiprocessing import context
import string
import time
from unicodedata import category
from dateutil.parser import parse
import numbers
from django.shortcuts import render, redirect, get_object_or_404
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.urls import reverse
from .forms import ProjectForm, UserForm, PostForm, RegisterForm
from .models import FileUploaded, Post, Project, User, Operation
from django.views.generic import TemplateView
from django.forms.models import model_to_dict
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import ensure_csrf_cookie
from slugify import slugify
from .util import convert_size,convert_type
import pandas as pd
import openpyxl
import os, math
from json import dumps
import numpy as np

import io
from django.http import FileResponse
# from reportlab.pdfgen import canvas


# Create your views here.


def index(request):
    return render(request, 'index.html', locals())


def login(request):
    print(request.session.get('is_login', None))
    if request.session.get('is_login', None):
        return redirect('/post_list/')

    if request.method == "POST":
        login_form = UserForm(request.POST)
        message = "请检查填写的内容！"
        if login_form.is_valid():
            email = login_form.cleaned_data['email']
            password = login_form.cleaned_data['password']
            try:
                user = User.objects.get(email=email)
                if user.password == password:
                    request.session['is_login'] = True
                    request.session['user_id'] = user.id
                    request.session['email'] = user.email
                    return redirect('dataapp:post_list')
                else:
                    message = "密码不正确！"
            except:
                message = "用户不存在！"
        return render(request, 'login.html', locals())

    login_form = UserForm()
    return render(request, 'login.html', locals())


def register(request):
    if request.session.get('is_login', None):
        # 登录状态不允许注册
        return redirect("dataapp:index")
    if request.method == "POST":
        register_form = RegisterForm(request.POST)
        print(request.POST)
        message = "请检查填写的内容！"
        if register_form.is_valid():  # 获取数据
            username = register_form.cleaned_data['username']
            password1 = register_form.cleaned_data['password1']
            password2 = register_form.cleaned_data['password2']
            email = register_form.cleaned_data['email']
            if password1 != password2:  # 判断两次密码是否相同
                message = "两次输入的密码不同！"
                return render(request, 'register.html', locals())
            else:
                same_name_user = User.objects.filter(username=username)
                if same_name_user:  # 用户名唯一
                    message = '用户已经存在，请重新选择用户名！'
                    return render(request, 'register.html', locals())
                same_email_user = User.objects.filter(email=email)
                if same_email_user:  # 邮箱地址唯一
                    message = '该邮箱地址已被注册，请使用别的邮箱！'
                    return render(request, 'register.html', locals())

                # 当一切都OK的情况下，创建新用户

                new_user = User.objects.create()
                new_user.username = username
                new_user.password = password1
                new_user.email = email
                new_user.save()
                return redirect('dataapp:login')  # 自动跳转到登录页面
    register_form = RegisterForm()
    return render(request, 'register.html', locals())


def logout(request):
    if not request.session.get('is_login', None):
        return redirect('dataapp:index')
    request.session.flush()

    return redirect('/')

def file_list(request,id):
    project = Project.objects.get(id=id)
    files = project.project_files.all()
    filelist = []
    print(files)
    for f in files:
        file_info={'id':f.id,
        'name':os.path.split(f.file.name)[1],
        'size':convert_size(f.file.size),
        'type':convert_type(f.file.name),
        'content':''}
        filelist.append(file_info)
        #os.stat(f.file).st_size
    return JsonResponse({'data':filelist})

def post_list(request):
    if not request.session.get('is_login', None):
        # 非登录状态显示主页
        return redirect("/")
    print('-------email-------',request.session['email'])
    user = User.objects.filter(email=request.session['email']).first()
    posts = user.user_posts.all().filter(category='table')
    # for post in posts:
    # print(post)
    # print(post.get_absolute_url())
    return render(request, 'post_list.html', {'posts': posts})

def text_list(request):
    if not request.session.get('is_login', None):
        # 非登录状态显示主页
        return redirect("/")
    user = User.objects.filter(email=request.session['email']).first()
    posts = user.user_posts.all().filter(category='text')
    # for post in posts:
    # print(post)
    # print(post.get_absolute_url())
    return render(request, 'text_list.html', {'posts': posts})



class MainView(TemplateView):
    template_name = 'dataapp/main.html'

@ensure_csrf_cookie
def get_csrftoken(request):
    return JsonResponse({'status':'ok'})

@ensure_csrf_cookie
def create_project_view(request):
    if request.method == "POST":
        project_form = ProjectForm(request.POST)
        if project_form.is_valid():
            current_project = Project.objects.filter(title=project_form.cleaned_data['title'])
            if current_project:
                #duplicate!
                return JsonResponse({'status':0, 'msg':'已经存在同名项目！'})
            # upload files
            print(request.FILES.getlist('files[]'))
            
            project_form.save()
            for f in request.FILES.getlist('files[]'):    
                instance = FileUploaded(project_name=project_form.clean().get('title'), file=f,project=project_form.instance)
                instance.save()
            
            return JsonResponse({'status':1, 'msg':'创建项目成功!'})
        else:
            project_form = ProjectForm()
            return JsonResponse({'status':0,'msg':'Invalid form data'})
    else:
        return HttpResponse()


def project_info(request,id):
    if request.method=='GET':
        
        project = Project.objects.get(id=id)
        print(project.created)
        pd = Project.to_dict(project)
        print(pd)
        return JsonResponse({'data':pd})
    else:
        # print('in')
        return HttpResponse()


def file_upload_view(request):
    print(request.FILES)
    print(request.POST)
    print(request.POST.get('title'))
    print(request.POST.get('body'))
    print(request.FILES.get('file[0]'))
    print(request.session['email'])

    if not request.session.get('is_login', None):
        # 非登录状态显示主页
        return redirect("/")

    if request.method == "POST":

        myfile = {'upload': request.FILES.get('file[0]')}
        the_user = User.objects.get(email=request.session['email'])
        my_request = request.POST.copy()
        my_request['author'] = the_user.id

        post_form = PostForm(my_request, myfile)
        if post_form.is_valid():
            title = post_form.cleaned_data['title']
            same_name_title = Post.objects.filter(title=title)
            if same_name_title:
                print('项目已经存在')
                message = '项目已经存在，请重新选择项目名称！'
                return render(request, 'new_post.html', locals())
                
            post_form.cleaned_data['slug'] = slugify(post_form.cleaned_data['title'])

            attachment = post_form.save()
            # 为新增的数据添加slug
            new_post = Post.objects.filter(title=post_form.cleaned_data['title']).get()
            new_post.slug = slugify(post_form.cleaned_data['title'])
            new_post.save()
            print('项目添加成功')

            # data = {'is_valid': True, 'url': attachment.upload.url, 'name': attachment.upload.name}
            # return redirect("dataapp:post_list")
            if post_form.cleaned_data['category']=='table':
                return render(request, 'post_list.html', locals())
            elif post_form.cleaned_data['category']=='text':
                return render(request, 'text_list.html', locals())
            elif post_form.cleaned_data['category']=='speech':
                return render(request, 'speech_list.html', locals())
            else:
                return render(request, 'image_list.html', locals())
    post_form = PostForm()
    if request.method == 'GET':
        return render(request, "new_post.html")


def post_detail(request, post):
    post = get_object_or_404(Post, slug=post)
    path = 'media/' + str(post.upload)
    all_data = pd.read_csv(path)
    number_per_page = 15
    data = all_data.sample(n=40)
    total = all_data.shape[0] * all_data.shape[1]
    print('数据总量',total)
    high_level_cols = ['name', 'phone_number']

    if request.GET.get('high_level_handle_time') == None:
        data = all_data.sample(n=40)

    # 点击高风险数据处理
    if request.GET.get('high_level_handle_time'):
        # operations = Operation.objects.all()
        # operations.delete()
        quantity_of_high_level = sum(all_data['name'].apply(lambda x: not str(x).count('*'))) + sum(
            [not str(i).count('*') for i in all_data['phone_number']])
        if quantity_of_high_level != 0:
            all_data.loc[[not str(i).count('*') for i in all_data['phone_number']], 'phone_number'] = all_data.loc[[
                                                                                                                       not str(
                                                                                                                           i).count(
                                                                                                                           '*')
                                                                                                                       for
                                                                                                                       i
                                                                                                                       in
                                                                                                                       all_data[
                                                                                                                           'phone_number']], 'phone_number'].astype(
                str).str[0:3] + '****' + all_data.loc[[not str(i).count('*') for i in
                                                       all_data['phone_number']], 'phone_number'].astype(str).str[7:]
            all_data.loc[all_data['name'].apply(lambda x: not str(x).count('*')), 'name'] = all_data.loc[
                                                                                                all_data['name'].apply(
                                                                                                    lambda x: not str(
                                                                                                        x).count(
                                                                                                        '*')), 'name'].astype(
                str).str[0:1] + '*' + all_data.loc[
                                          all_data['name'].apply(lambda x: not str(x).count('*')), 'name'].astype(
                str).str[2:]

            save_path = os.path.splitext(path)[0] + '.csv'
            all_data.to_csv(save_path, index=False, sep=',', encoding="utf_8_sig")
            data = all_data.sample(n=40)
            new_operation = Operation.objects.create()
            print('高风险数据量', quantity_of_high_level)
            new_operation.content = '您成功处理了包括‘姓名’、‘电话’等高风险敏感信息'
            new_operation.created = parse(request.GET.get('high_level_handle_time'))
            new_operation.task = post
            new_operation.quantity = quantity_of_high_level
            new_operation.high_proportion = quantity_of_high_level/total
            new_operation.save()
            print('高风险数处理时间', request.GET.get('high_level_handle_time'))

    # 点击处理中风险数据
    if request.GET.get('middle_level_handle_time'):
        # operations = Operation.objects.all()
        # operations.delete()
        quantity_of_middle_level = sum(all_data['address'].apply(lambda x: not str(x).count('*')))
        if quantity_of_middle_level != 0:
            all_data.loc[all_data['address'].apply(lambda x: not str(x).count('*')), 'address'] = all_data.loc[all_data[
                                                                                                                   'address'].apply(
                lambda x: not str(x).count('*')), 'address'].astype(str).str[0:-13] + '**********'
            save_path = os.path.splitext(path)[0] + '.csv'
            all_data.to_csv(save_path, index=False, sep=',', encoding="utf_8_sig")
            data = all_data.sample(n=40)
            new_operation = Operation.objects.create()
            print('中风险数据量', quantity_of_middle_level)
            new_operation.content = '您成功处理了包括‘住址’、‘IP地址’等中风险敏感信息'
            new_operation.created = parse(request.GET.get('middle_level_handle_time'))
            new_operation.task = post
            new_operation.quantity = quantity_of_middle_level
            new_operation.middle_proportion = quantity_of_middle_level/total
            new_operation.save()
            print('中风险数处理时间', request.GET.get('middle_level_handle_time'))

    if request.GET.get('low_level_handle_time'):
        # operations = Operation.objects.all()
        # operations.delete()
        quantity_of_low_level = sum(all_data['age'].apply(lambda x: str(x) != 'nan'))
        if quantity_of_low_level != 0:
            all_data.loc[all_data['age'].apply(lambda x: str(x) != 'nan'), 'age'] = 'nan'
            save_path = os.path.splitext(path)[0] + '.csv'
            all_data.to_csv(save_path, index=False, sep=',', encoding="utf_8_sig")
            data = all_data.sample(n=40)
            new_operation = Operation.objects.create()
            print('低风险数据量', quantity_of_low_level)
            new_operation.content = '您成功处理了包括‘年龄’、‘统计数据’等低风险敏感信息'
            new_operation.created = parse(request.GET.get('low_level_handle_time'))
            new_operation.task = post
            new_operation.quantity = quantity_of_low_level
            new_operation.low_proportion  = quantity_of_low_level/total
            new_operation.save()
            print('低风险数处理时间', request.GET.get('low_level_handle_time'))

    # if os.path.splitext(path)[-1]=='.csv':
    # 将数据转换成列表
    table_head = list(data.columns)
    excel_data = list()
    for row in range(data.shape[0]):
        excel_data.append(list(data.iloc[row]))

    # 将数据分页展示
    paginator = Paginator(excel_data, number_per_page)  # 每页15条数据
    page = request.GET.get('page')

    try:
        excel_data = paginator.page(page)
    except PageNotAnInteger:
        excel_data = paginator.page(1)
    except EmptyPage:
        excel_data = paginator.page(paginator.num_pages)

    # 找出高风险数据
    step = number_per_page
    data_slice = [data[i:i + step] for i in range(0, len(data), step)]  # 分页查找
    high_level = {}
    for p in range(len(data_slice) + 1):
        if page == None:
            page = 1
        if p == int(page) - 1:
            cols = list(data_slice[p])
            high_level_cols = ['name', 'phone_number']
            # data_slice[p]['phone_number']=data_slice[p]['phone_number'].astype(str)
            for col in high_level_cols:
                if col in cols:
                    for item in list(data_slice[p][col]):
                        if not item.count('*'):
                            if not list(data_slice[p][col]).index(item) in high_level.keys():
                                high_level[list(data_slice[p][col]).index(item)] = [cols.index(col), ]
                            else:
                                high_level[list(data_slice[p][col]).index(item)].append(cols.index(col))

    # 找出中风险数据
    middle_level = {}
    for p in range(len(data_slice) + 1):
        if page == None:
            page = 1
        if p == int(page) - 1:
            cols = list(data_slice[p])
            middle_level_cols = ['address']
            for col in middle_level_cols:
                if col in cols:
                    b = -1
                    for item in list(data_slice[p][col]):
                        if not item.count('*'):
                            b = list(data_slice[p][col]).index(item, b + 1, len(list(data_slice[p][col])))
                            if b not in middle_level.keys():
                                middle_level[b] = [cols.index(col), ]
                            else:
                                middle_level[b].append(cols.index(col))

    # 找出低风险数据
    low_level = {}
    for p in range(len(data_slice) + 1):
        if page == None:
            page = 1
        if p == int(page) - 1:
            cols = list(data_slice[p])
            low_level_cols = ['age']
            for col in low_level_cols:
                if col in cols:
                    b = -1
                    for item in list(data_slice[p][col]):
                        if str(item) != 'nan':
                            b = list(data_slice[p][col]).index(item, b + 1, len(list(data_slice[p][col])))
                            if b not in low_level.keys():
                                low_level[b] = [cols.index(col), ]
                            else:
                                low_level[b].append(cols.index(col))

    dataJSON = dumps(high_level)
    dataJSON_middle = dumps(middle_level)
    dataJSON_low = dumps(low_level)

    print('高风险', high_level)
    print('中风险', middle_level)
    print('低风险', low_level)

    return render(request, 'post_detail.html', {'post': post, 'excel_data': excel_data,
                                                page: 'pages', 'data': dataJSON, 'middle_data': dataJSON_middle,
                                                'low_data': dataJSON_low, 'table_head': table_head})


def handle_high_level(request, post):
    post = get_object_or_404(Post, slug=post)
    path = 'media/' + str(post.upload)
    all_data = pd.read_csv(path)
    all_data.loc[[not str(i).count('*') for i in all_data['phone_number']], 'phone_number'] = all_data.loc[
                                                                                                  [not str(i).count('*')
                                                                                                   for i in all_data[
                                                                                                       'phone_number']], 'phone_number'].astype(
        str).str[0:3] + '****' + all_data.loc[
                                     [not str(i).count('*') for i in all_data['phone_number']], 'phone_number'].astype(
        str).str[7:]
    data = all_data.sample(n=40)
    number_per_page = 15

    # if os.path.splitext(path)[-1]=='.csv':
    
    table_head = list(data.columns)
    excel_data = list()
    for row in range(data.shape[0]):
        excel_data.append(list(data.iloc[row]))

    # 将数据分页展示
    paginator = Paginator(excel_data, number_per_page)  # 每页15条数据
    page = request.GET.get('page')

    try:
        excel_data = paginator.page(page)
    except PageNotAnInteger:
        excel_data = paginator.page(1)
    except EmptyPage:
        excel_data = paginator.page(paginator.num_pages)

    return render(request, 'post_detail.html',
                  {'post': post, 'excel_data': excel_data, page: 'pages', 'table_head': table_head})

@ensure_csrf_cookie
def project_list(request):
    # print(request.POST.get('category'))
    if request.method == 'POST':
        project_list = Project.objects.filter(category = request.POST.get('category')).values()
        #print(project_list)
        plist = list(project_list)
        print(plist)
        return JsonResponse({'data':plist,'status':1})
    else:
        return JsonResponse({'msg':'Invalid method','status':0})

@ensure_csrf_cookie
def project_delete(request):
    project = Project.objects.get(id=request.POST.get('id'))
    project.delete()
    path = 'media/project_'+str(project.title)
    print(path)
    os.removedirs(path)
    return HttpResponse()

def post_delete(requst, pk):
    post = Post.objects.get(id=pk)
    post.delete()
    path = 'media/'+str(post.upload)
    print(path)
    os.remove(path)
    if post.category=='table':
        return redirect("dataapp:post_list")
    if post.category=='text':
        return redirect("dataapp:text_list")

# table数据生成报告
def generate_report(request, post):
    post = get_object_or_404(Post, slug=post)
    operation = post.post_operation.all()
    print(len(operation))
    generate_report_time = time.strftime('%Y/%m/%d', time.localtime(time.time()))
    if len(operation) != 0:
        return render(request, "generate_report.html", {'post': post, 'operation': operation,
                                             'generate_report_time': generate_report_time})
    if post.category=='table':
        return redirect("dataapp:post_detail", post=post.slug)
    if post.category=='text':
        return redirect("dataapp:text_detail", post=post.slug)


class DFAFilter(object):
    def __init__(self):
        self.keyword_chains = {}  # 敏感词链表
        self.delimit = '\x00'  # 敏感词词尾标识
        self.whitelist=['的'] # 白名单
        self.uncertain=['性','日','死'] #不确定的，多为单个字
    def add(self, keyword):
        flag = 1
        if keyword in self.uncertain:
            flag = 2
        keyword = keyword.lower()  # 关键词英文变为小写
        chars = keyword.strip()  # 关键字去除首尾空格和换行
        if not chars:  # 如果关键词为空直接返回
            return
        level = self.keyword_chains
        # 遍历关键字的每个字
        for i in range(len(chars)):
            # 如果这个字已经存在字符链的key中就进入其子字典
            if chars[i] in level:
                level = level[chars[i]]
            else:
                if not isinstance(level, dict):
                    break
                for j in range(i, len(chars)):
                    level[chars[j]] = {}
                    last_level, last_char = level, chars[j]
                    level = level[chars[j]]
                last_level[last_char] = {self.delimit: flag}
                break
        if i == len(chars) - 1:
            level[self.delimit] = flag
        
    def init_chains(self, path):
        import docx
        import re
        keyword_file = docx.Document(path)
        keywords = []
        for para in keyword_file.paragraphs:
            keywords.extend(re.split("[；|;]",str(para.text)))
        # print(keywords)
        for keyword in keywords:
            self.add(str(keyword).strip())

    def filter(self, message):
        message = message.lower()
        ret = []  # 返回字符串列表，为了便于前端显示，采取形如[{'flag':0, 'text':"abc"}, {'flag':1 ,'text':"sb"}]的形式返回
        start = 0
        while start < len(message):
            level = self.keyword_chains
            step_ins = 0
            for char in message[start:]:
                if char in level:
                    step_ins += 1
                    if self.delimit not in level[char] or message[start:start+step_ins] in self.whitelist: # current serial is legal
                        level = level[char]
                    else:
                        ret.append({'flag':level[char][self.delimit],'text':message[start:start+step_ins]})
                        start += step_ins - 1
                        break
                else:
                    ret.append({'flag':0,'text':message[start]})
                    break
            else:
                ret.append({'flag':0,'text':message[start]})
            start += 1

        return ret

def search_keyword(request):
    import re
    import docx
    path = 'media/filter/keywords.docx'
    file_path = 'media/files/text.docx'
    content = []
    content_str =''
    file = docx.Document(file_path)
    for para in file.paragraphs:
        content_str+=str(para.text)
    # content_str = '卧槽你妈的，你是傻逼吗性爱电影，狗东西是个什么玩意儿李红智'
    filter = DFAFilter()
    filter.init_chains(path)
    res = filter.filter(content_str)
    print(res)
    return JsonResponse({'data':res})

def text_censor(request,fid):
    import re
    import docx
    path = 'media/filter/keywords.docx'
    file_path = FileUploaded.objects.get(id=fid).file.path
    content = []
    content_str =''
    file = docx.Document(file_path)
    for para in file.paragraphs:
        content_str+=str(para.text)
    filter = DFAFilter()
    filter.init_chains(path)
    res = filter.filter(content_str)
    print(res)
    return JsonResponse({'data':res})


def text_detail(request, post):
    import re
    import docx
    post = get_object_or_404(Post, slug=post)
    path = 'media/' + str(post.upload)
    content=[]
    content_string = ''
    file = docx.Document(path)
    
    for para in file.paragraphs:
        content.append(str(para.text))
        content_string = content_string+str(para.text)

    print(len(content_string))

    phoneRegex = re.compile(r'''(
               (13\d|14[5|7]|15\d|166|17[3|6|7]|18\d)\d{8}
                )''', re.VERBOSE)


    idnumberRegex = re.compile(r'''(
               [1-9]\d{5}[12]\d{3}(0[1-9]|1[012])(0[1-9]|[12][0-9]|3[01])\d{3}[0-9xX]
                )''', re.VERBOSE)
    
    emailRegex = re.compile(r'''(
                [a-zA-Z0-9._%+-]+          
                @
                [a-zA-Z0-9.-]+             
                (\.[a-zA-Z]{2,4})
                )''', re.VERBOSE)

    adressRegex = re.compile(r'''(
                ([\u4e00-\u9fa5]{2}?(?:市)){0,1}
                ([\u4e00-\u9fa5]{2,5}?(?:区|县|州)){1}
                ([\u4e00-\u9fa5]{2,5}?(?:村|镇|街道|路)){1}
                ([\d]{1,3}(?:号)){1}
                ([\u4e00-\u9fa5]{2,5}?(?:小区)){0,1}
                ([\d]{1,3}(?:号楼|幢)){0,1}
                ([\d]{1,3}(?:单元)){0,1}
                ([\d]{1,5}(?:室){0,1}){0,1}
                )''', re.VERBOSE)
        
    high_level = {'number':[],'name':[]}
    middle_level = []
    low_level = []
    # print(content_string)

    for groups in phoneRegex.findall(content_string) :
        high_level['number'].append(groups[0])

    for groups in idnumberRegex.findall(content_string) :
        high_level['number'].append(groups[0])

    for groups in emailRegex.findall(content_string) :
        middle_level.append(groups[0])

    for groups in adressRegex.findall(content_string) :
        if len(groups[0]):
            middle_level.append(groups[0])


    thename=bert.name_predict(input_str=content_string)
    if len(thename):
        high_level['name']=high_level['name']+thename

    


    name_compliance = [x[0]+'*'*(len(x)-1) for x in high_level['name']]
    number_compliance = [x[0:3]+'****'+x[7:] if len(x)==11 else x[0:6]+'********'+x[14:] for x in high_level['number']]
    high_level_compliance = number_compliance+name_compliance
    middle_level_compliance = middle_level.copy()

    middleRegex = re.compile(r'''(              
                ([\d]{1,5})
                )''', re.VERBOSE)
    for x in range(len(middle_level_compliance)):
        for groups in middleRegex.findall(middle_level_compliance[x]):
            middle_level_compliance[x]=middle_level_compliance[x].replace(groups[0],'*'*len(groups[0]))

    

    data_to_frontend = {'high_level':high_level['number']+high_level['name'],'middle_level':middle_level}
    dataJSON = dumps(data_to_frontend)



    high_level_dict=dict(zip(high_level['number']+high_level['name'], high_level_compliance))
    print('高风险',high_level_dict)

    # 点击高风险数据处理（电话，身份证号）
    if request.GET.get('high_level_handle_time'):
        quantity_of_high_level = len(high_level['number']+high_level['name'])
        if quantity_of_high_level != 0:
            for k,v in high_level_dict.items():
                for index,para in enumerate(content):
                    if k in para:
                        content[index]=content[index].replace(k,v)

            
            save_path = os.path.splitext(path)[0] + '.docx'
            doc = docx.Document()
            for para in content:
                par = doc.add_paragraph('')
                run_ = par.add_run(para)
            doc.save(save_path)

            # all_data.to_csv(save_path, index=False, sep=',', encoding="utf_8_sig")
            # data = all_data.sample(n=40)
            new_operation = Operation.objects.create()
            print('高风险数据量', quantity_of_high_level)
            new_operation.content = '您成功处理了包括‘姓名’、‘电话’等高风险敏感信息'
            new_operation.created = parse(request.GET.get('high_level_handle_time'))
            new_operation.task = post
            new_operation.quantity = quantity_of_high_level
            new_operation.high_proportion = quantity_of_high_level/len(content_string)
            new_operation.save()
            print('高风险数处理时间', request.GET.get('high_level_handle_time'))



    middle_level_dict=dict(zip(middle_level, middle_level_compliance))
    print('中风险',middle_level_dict)
   

    # 点击中风险数据处理（地理位置，邮件）
    if request.GET.get('middle_level_handle_time'):
        quantity_of_middle_level = len(middle_level)
        if quantity_of_middle_level != 0:
            for k,v in middle_level_dict.items():
                for index,para in enumerate(content):
                    if k in para:
                        content[index]=content[index].replace(k,v)

            
            save_path = os.path.splitext(path)[0] + '.docx'
            doc = docx.Document()
            for para in content:
                par = doc.add_paragraph('')
                run_ = par.add_run(para)
            doc.save(save_path)

            # all_data.to_csv(save_path, index=False, sep=',', encoding="utf_8_sig")
            # data = all_data.sample(n=40)
            new_operation = Operation.objects.create()
            print('中风险数据量', quantity_of_middle_level)
            new_operation.content = '您成功处理了包括‘地址’、等中风险敏感信息'
            new_operation.created = parse(request.GET.get('middle_level_handle_time'))
            new_operation.task = post
            new_operation.quantity = quantity_of_middle_level
            new_operation.high_proportion = quantity_of_middle_level/len(content_string)
            new_operation.save()
            print('中风险数处理时间', request.GET.get('middle_level_handle_time'))
       
    return render(request, 'text_detail.html', {'post': post, 'content':content ,'data_to_frontend':dataJSON})

def test_db(request):
    address = request.POST.get('address')
    user = request.POST.get('user')
    pwd = request.POST.get('pwd')
    dbname = request.POST.get('dbname')
    conn = DBConnection(user,pwd,address,dbname)
    ret = conn.test_fetch()
    return JsonResponse({"data":ret})




def image_detail(request):
    return render(request, 'image_detail.html') 

def speech_detail(request):
    return render(request, 'speech_detail.html') 

def project_configration(request):
    return render(request, 'project_configration.html') 




